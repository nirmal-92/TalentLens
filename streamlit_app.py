import streamlit as st
from groq import Groq
from pypdf import PdfReader
from docx import Document
import streamlit.components.v1 as components
import io, json, re, time, uuid
from datetime import datetime
import numpy as np
import plotly.graph_objects as go

# ── PAGE CONFIG ────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="TalentLens · AI Hiring Suite",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ── HARDCODED API KEY (Backend only — never shown in UI) ───────────────────────
GROQ_API_KEY = st.secrets["GROQ_API_KEY"]   

# ── CUSTOM CSS ─────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Sora:wght@300;400;600;700&family=JetBrains+Mono:wght@400;600&display=swap');

html, body, [class*="css"] { font-family: 'Sora', sans-serif; }
.stApp { background: #0a0e1a; color: #e8eaf0; }

[data-testid="stSidebar"] { background: #0f1628 !important; border-right: 1px solid #1e2a45; }
[data-testid="stSidebar"] * { color: #c8cfe0 !important; }

.hero {
    background: linear-gradient(135deg, #0f1628 0%, #1a2540 50%, #0f1628 100%);
    border: 1px solid #1e3a6e; border-radius: 16px;
    padding: 40px 48px; margin-bottom: 32px; position: relative; overflow: hidden;
}
.hero::before {
    content: ''; position: absolute; top: -50%; right: -10%;
    width: 400px; height: 400px;
    background: radial-gradient(circle, rgba(59,130,246,0.12) 0%, transparent 70%);
    border-radius: 50%;
}
.hero-title {
    font-size: 2.8rem; font-weight: 700;
    background: linear-gradient(90deg, #60a5fa, #a78bfa, #34d399);
    -webkit-background-clip: text; -webkit-text-fill-color: transparent;
    margin: 0 0 8px 0; letter-spacing: -0.5px;
}
.hero-sub { font-size: 1.05rem; color: #7090b0; margin: 0; font-weight: 300; }

.metric-row { display: flex; gap: 16px; margin-bottom: 28px; flex-wrap: wrap; }
.metric-card {
    background: #111827; border: 1px solid #1e2a45; border-radius: 12px;
    padding: 20px 24px; flex: 1; min-width: 140px; text-align: center;
}
.metric-num { font-size: 2rem; font-weight: 700; color: #60a5fa; font-family: 'JetBrains Mono', monospace; }
.metric-label { font-size: 0.78rem; color: #5a7090; text-transform: uppercase; letter-spacing: 1px; margin-top: 4px; }

.score-pill { display: inline-block; padding: 4px 14px; border-radius: 999px; font-family: 'JetBrains Mono', monospace; font-weight: 600; font-size: 0.9rem; }
.score-high { background: rgba(52,211,153,0.15); color: #34d399; border: 1px solid rgba(52,211,153,0.3); }
.score-mid  { background: rgba(251,191,36,0.15);  color: #fbbf24; border: 1px solid rgba(251,191,36,0.3); }
.score-low  { background: rgba(248,113,113,0.15); color: #f87171; border: 1px solid rgba(248,113,113,0.3); }

.candidate-card {
    background: #111827; border: 1px solid #1e2a45; border-radius: 14px;
    padding: 24px; margin-bottom: 16px;
}
.candidate-name { font-size: 1.2rem; font-weight: 600; color: #e8eaf0; }
.candidate-meta { font-size: 0.82rem; color: #5a7090; margin-top: 4px; }

.skill-tag {
    display: inline-block; background: rgba(59,130,246,0.1); color: #60a5fa;
    border: 1px solid rgba(59,130,246,0.25); border-radius: 6px;
    padding: 2px 10px; font-size: 0.78rem; margin: 2px;
}
.project-tag {
    display: inline-block; background: rgba(167,139,250,0.1); color: #a78bfa;
    border: 1px solid rgba(167,139,250,0.25); border-radius: 6px;
    padding: 2px 10px; font-size: 0.78rem; margin: 2px;
}
.intern-tag {
    display: inline-block; background: rgba(52,211,153,0.1); color: #34d399;
    border: 1px solid rgba(52,211,153,0.25); border-radius: 6px;
    padding: 2px 10px; font-size: 0.78rem; margin: 2px;
}

.bar-wrap { background: #1a2235; border-radius: 999px; height: 8px; margin: 6px 0 14px 0; }
.bar-fill  { height: 8px; border-radius: 999px; }

.section-title {
    font-size: 1.4rem; font-weight: 600; color: #e8eaf0;
    margin: 32px 0 16px 0; padding-bottom: 10px; border-bottom: 1px solid #1e2a45;
}

.rec-badge { display: inline-block; padding: 6px 16px; border-radius: 8px; font-size: 0.85rem; font-weight: 600; }
.rec-hire  { background: rgba(52,211,153,0.15); color: #34d399; border: 1px solid rgba(52,211,153,0.4); }
.rec-maybe { background: rgba(251,191,36,0.15);  color: #fbbf24; border: 1px solid rgba(251,191,36,0.4); }
.rec-nope  { background: rgba(248,113,113,0.15); color: #f87171; border: 1px solid rgba(248,113,113,0.4); }

.jd-box {
    background: #0d1422; border: 1px solid #1e3050; border-radius: 12px;
    padding: 24px; font-size: 0.92rem; line-height: 1.7; color: #b0c4de;
    white-space: pre-wrap; max-height: 500px; overflow-y: auto;
}

.interview-card {
    background: #0d1422; border: 1px solid #1e3a6e; border-radius: 12px;
    padding: 20px; margin-bottom: 16px;
}
.interview-q { color: #60a5fa; font-weight: 600; font-size: 0.95rem; margin-bottom: 8px; }
.interview-a { color: #b0c4de; font-size: 0.88rem; line-height: 1.6; }
.interview-flag { color: #f87171; font-size: 0.82rem; margin-top: 6px; }

/* Interview AI styles */
.ai-card{background:#111827;border:1px solid #1e2a45;border-radius:16px;padding:1.5rem;margin-bottom:1rem;}
.ai-card-title{font-family:'JetBrains Mono',monospace;font-size:10px;letter-spacing:.12em;text-transform:uppercase;color:#0df2c8;margin-bottom:12px;}
.q-bubble{background:#0d1422;border:1px solid #1e2a45;border-radius:12px;padding:16px 18px;font-size:15px;line-height:1.65;color:#e8eaf0;margin-bottom:1rem;}
.answer-bubble{background:rgba(13,242,200,.07);border:1px solid rgba(13,242,200,.25);border-radius:12px;padding:12px 16px;font-size:13px;line-height:1.6;color:#7090b0;margin-top:8px;border-left:3px solid #09b394;}
.feedback-bubble{font-size:12px;font-style:italic;color:#7090b0;padding:8px 12px;border-radius:8px;background:rgba(77,166,255,.05);border:1px solid rgba(77,166,255,.15);margin-top:8px;}
.verdict-box{background:linear-gradient(135deg,rgba(13,242,200,.04),rgba(77,166,255,.04));border:1px solid #2a5a8a;border-radius:14px;padding:1.5rem;margin-top:1rem;}
.tx-pill{display:inline-flex;align-items:center;gap:6px;background:rgba(13,242,200,.08);border:1px solid #09b394;
  border-radius:20px;padding:3px 10px;font-family:'JetBrains Mono',monospace;font-size:9px;color:#0df2c8;margin-bottom:6px;}
.flag-green{background:rgba(46,204,130,.1);border:1px solid #2ecc82;border-radius:8px;padding:5px 10px;
  font-size:12px;color:#2ecc82;margin:3px 0;display:block;}
.flag-red{background:rgba(255,78,106,.1);border:1px solid #ff4e6a;border-radius:8px;padding:5px 10px;
  font-size:12px;color:#ff4e6a;margin:3px 0;display:block;}

.tech-tag-auto {
    display: inline-block; background: rgba(13,242,200,0.1); color: #0df2c8;
    border: 1px solid rgba(13,242,200,0.3); border-radius: 6px;
    padding: 3px 10px; font-size: 0.78rem; margin: 2px;
    font-family: 'JetBrains Mono', monospace;
}

.top-resume-banner {
    background: linear-gradient(135deg, rgba(52,211,153,0.08), rgba(13,242,200,0.05));
    border: 1px solid rgba(52,211,153,0.35); border-radius: 12px;
    padding: 16px 20px; margin-bottom: 16px;
}

.stTextInput input, .stTextArea textarea {
    background: #111827 !important; border: 1px solid #1e2a45 !important;
    color: #e8eaf0 !important; border-radius: 10px !important;
}
.stButton > button {
    background: linear-gradient(135deg, #2563eb, #7c3aed) !important;
    color: white !important; border: none !important; border-radius: 10px !important;
    font-family: 'Sora', sans-serif !important; font-weight: 600 !important;
    padding: 10px 28px !important;
}
.stProgress > div > div { background: linear-gradient(90deg, #2563eb, #7c3aed) !important; }
details { background: #111827 !important; border: 1px solid #1e2a45 !important; border-radius: 12px !important; }
summary { color: #e8eaf0 !important; font-weight: 600 !important; }
hr { border-color: #1e2a45 !important; }
.stSlider > div > div > div { background: #2563eb !important; }
</style>
""", unsafe_allow_html=True)

# ── SESSION STATE ──────────────────────────────────────────────────────────────
for k, v in {
    "parsed_resumes": [],
    "jd_text": "",
    "jd_requirement": "",
    "jd_tech_stack": "",
    "scored_results": [],
    "interview_results": {},
    "groq_client": None,
    "weights": {"skills": 30, "experience": 30, "education": 20, "culture_fit": 20},
    # AI Interview state
    "iv_page": "setup",
    "iv_session_id": "",
    "iv_questions": [],
    "iv_answers": [],
    "iv_q_idx": 0,
    "iv_resume_text": "",
    "iv_report": None,
    "iv_q_start_time": None,
    "iv_live_transcript": "",
    "iv_typed_answer": "",
    "iv_show_edit": False,
    "iv_candidate_name": "",
    "iv_auto_loaded": False,   # tracks whether top resume was auto-loaded
}.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ── GROQ — Initialize from hardcoded key ──────────────────────────────────────
if st.session_state.groq_client is None:
    try:
        st.session_state.groq_client = Groq(api_key=GROQ_API_KEY)
    except Exception as e:
        st.error(f"❌ Could not initialize Groq client: {e}")

GROQ_MODEL = "llama-3.3-70b-versatile"

def call_groq(client, prompt, retries=3, max_tokens=3000):
    c = client or st.session_state.groq_client
    for attempt in range(retries):
        try:
            resp = c.chat.completions.create(
                model=GROQ_MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
                max_tokens=max_tokens,
            )
            return resp.choices[0].message.content
        except Exception as e:
            if "429" in str(e) or "rate_limit" in str(e).lower():
                wait = 12 * (attempt + 1)
                st.toast(f"Rate limit — waiting {wait}s...", icon="⏳")
                time.sleep(wait)
            else:
                raise e
    raise Exception("Groq quota exhausted. Wait a moment and retry.")

# ── FILE READERS ───────────────────────────────────────────────────────────────
def read_pdf(file):
    try:
        raw = file.read()
        reader = PdfReader(io.BytesIO(raw))
        return "\n".join(p.extract_text() or "" for p in reader.pages).strip()
    except Exception:
        return ""

def read_docx(file):
    try:
        doc = Document(io.BytesIO(file.read()))
        return "\n".join(p.text for p in doc.paragraphs).strip()
    except Exception:
        return ""

def extract_text_bytes(b, filename):
    fn = filename.lower()
    if fn.endswith(".pdf"):
        try:
            reader = PdfReader(io.BytesIO(b))
            return "\n".join(p.extract_text() or "" for p in reader.pages).strip()
        except Exception as e:
            return f"[PDF error: {e}]"
    if fn.endswith(".docx"):
        try:
            doc = Document(io.BytesIO(b))
            return "\n".join(p.text for p in doc.paragraphs).strip()
        except Exception as e:
            return f"[DOCX error: {e}]"
    if fn.endswith(".txt"):
        return b.decode("utf-8", errors="replace")
    return "[Unsupported]"

# ── JSON HELPERS ───────────────────────────────────────────────────────────────
def clean_json(raw):
    raw = raw.strip()
    raw = re.sub(r"^```(?:json)?", "", raw)
    raw = re.sub(r"```$", "", raw)
    raw = raw.strip()
    arr_start = raw.find('[')
    obj_start = raw.find('{')
    if arr_start == -1 and obj_start == -1:
        return raw
    if arr_start == -1:
        start_char, start = '{', obj_start
    elif obj_start == -1:
        start_char, start = '[', arr_start
    else:
        if arr_start < obj_start:
            start_char, start = '[', arr_start
        else:
            start_char, start = '{', obj_start
    close_char = ']' if start_char == '[' else '}'
    depth = 0; end = -1; in_string = False; escape = False
    for i, ch in enumerate(raw[start:], start):
        if escape:
            escape = False; continue
        if ch == '\\' and in_string:
            escape = True; continue
        if ch == '"':
            in_string = not in_string; continue
        if in_string:
            continue
        if ch == start_char:
            depth += 1
        elif ch == close_char:
            depth -= 1
            if depth == 0:
                end = i + 1; break
    if end != -1:
        raw = raw[start:end]
    raw = re.sub(r',\s*([}\]])', r'\1', raw)
    return raw

def safe_int(val, default=0):
    try:
        return max(0, min(100, int(val)))
    except:
        return default

def ensure_list(data):
    if isinstance(data, list):
        return data
    if isinstance(data, dict):
        for key in ("questions", "data", "items", "interview_questions"):
            if key in data and isinstance(data[key], list):
                return data[key]
        if "question" in data:
            return [data]
    return []

# ── AI: EXTRACT TECH STACK FROM JD ────────────────────────────────────────────
def extract_tech_from_jd(client, jd_text):
    prompt = f"""Extract the technical skills, programming languages, frameworks, tools, and technologies mentioned in this Job Description.
Return ONLY a comma-separated list of technologies. No explanation, no JSON, no bullet points.
Example output: Python, FastAPI, PostgreSQL, Docker, AWS, React, Node.js

Job Description:
{jd_text[:3000]}"""
    raw = call_groq(client, prompt, max_tokens=200)
    return raw.strip().strip('"').strip("'")

# ── AI: PARSE RESUME ──────────────────────────────────────────────────────────
RESUME_SCHEMA = """{
  "name": "Full Name",
  "email": "email or empty",
  "phone": "phone or empty",
  "experience_years": 2,
  "current_title": "latest job title or empty",
  "skills": ["skill1", "skill2"],
  "technologies": ["tech1", "tech2"],
  "education": ["B.Tech Computer Science, Anna University 2023"],
  "certifications": ["AWS Certified Developer"],
  "internships": [
    {"company": "TCS", "role": "Software Intern", "duration": "3 months", "description": "Worked on REST APIs"}
  ],
  "projects": [
    {"name": "E-commerce Website", "tech_used": ["React","Node.js"], "description": "Built a full stack app with payment integration"}
  ],
  "key_achievements": ["Reduced API latency by 40%"],
  "summary": "2-3 sentence professional summary"
}"""

def parse_resume(client, text, filename):
    if not text or len(text.strip()) < 50:
        return _empty_resume(filename)
    prompt = f"""You are an expert resume parser. Extract ALL information from the resume carefully.
Return ONLY raw valid JSON. No markdown, no code fences, no explanation.
Required JSON structure:
{RESUME_SCHEMA}
Resume Text:
{text[:5000]}"""
    raw = call_groq(client, prompt)
    try:
        clean = re.sub(r"```json|```", "", raw).strip()
        match = re.search(r'\{.*\}', clean, re.DOTALL)
        data = json.loads(match.group() if match else clean)
    except Exception:
        data = _empty_resume(filename)
        data["summary"] = "Parsing failed — check file format."
    data["filename"] = filename
    data["raw_text"] = text
    return data

def _empty_resume(filename):
    return {
        "name": filename.replace(".pdf","").replace(".docx",""),
        "email": "", "phone": "", "experience_years": 0,
        "current_title": "", "skills": [], "technologies": [],
        "education": [], "certifications": [], "internships": [],
        "projects": [], "key_achievements": [],
        "summary": "Could not extract text from this file.",
        "filename": filename, "raw_text": ""
    }

# ── AI: GENERATE JD ───────────────────────────────────────────────────────────
def generate_jd(client, requirement):
    prompt = f"""You are an expert HR consultant. Write a professional Job Description.
Role: "{requirement}"
Sections: 1. Job Title 2. About the Role 3. Key Responsibilities (6-8 points)
4. Required Skills and Experience 5. Nice to Have 6. What We Offer
Be specific and professional."""
    return call_groq(client, prompt)

# ── AI: SCORE RESUME ──────────────────────────────────────────────────────────
SCORE_SCHEMA = """{
  "skills_score": 85, "experience_score": 80, "education_score": 75, "culture_fit_score": 78,
  "internship_score": 70, "project_score": 80,
  "strengths": ["strength1", "strength2", "strength3"],
  "gaps": ["gap1", "gap2"],
  "recommendation": "Strong Hire",
  "recommendation_reason": "One sentence explanation",
  "interview_questions": ["question1", "question2", "question3"],
  "verdict": "Recommended"
}"""

def score_resume(client, resume, jd, weights):
    internships_text = "".join(
        f"\n  - {i.get('role','')} at {i.get('company','')} ({i.get('duration','')}): {i.get('description','')}"
        for i in resume.get("internships", []))
    projects_text = "".join(
        f"\n  - {p.get('name','')}: {p.get('description','')} [Tech: {', '.join(p.get('tech_used', []))}]"
        for p in resume.get("projects", []))
    prompt = f"""You are a senior technical recruiter. Score this candidate against the job description.
Return ONLY raw valid JSON. No markdown, no code fences.
Scoring weights: Skills {weights['skills']}% | Experience {weights['experience']}% | Education {weights['education']}% | Culture Fit {weights['culture_fit']}%
recommendation must be one of: "Strong Hire","Hire","Maybe","Pass"
verdict must be one of: "Recommended","Consider","Not Recommended"
Required JSON: {SCORE_SCHEMA}
JOB DESCRIPTION: {jd[:2000]}
CANDIDATE: Name: {resume.get('name')} | Experience: {resume.get('experience_years')} years
Skills: {', '.join(resume.get('skills', []))} | Technologies: {', '.join(resume.get('technologies', []))}
Education: {', '.join(resume.get('education', []))} | Certifications: {', '.join(resume.get('certifications', []))}
Internships: {internships_text or 'None'} | Projects: {projects_text or 'None'}
Summary: {resume.get('summary')}"""
    raw = call_groq(client, prompt)
    try:
        clean = re.sub(r"```json|```", "", raw).strip()
        match = re.search(r'\{.*\}', clean, re.DOTALL)
        result = json.loads(match.group() if match else clean)
    except Exception:
        result = {"skills_score":0,"experience_score":0,"education_score":0,"culture_fit_score":0,
                  "internship_score":0,"project_score":0,"strengths":[],"gaps":["Scoring failed"],
                  "recommendation":"Maybe","recommendation_reason":"Could not evaluate",
                  "interview_questions":[],"verdict":"Consider"}
    ws = weights['skills']/100; we = weights['experience']/100
    wed = weights['education']/100; wc = weights['culture_fit']/100
    total_w = ws+we+wed+wc
    overall = int((result.get("skills_score",0)*ws + result.get("experience_score",0)*we +
                   result.get("education_score",0)*wed + result.get("culture_fit_score",0)*wc)/total_w)
    result["overall_score"] = overall
    result["name"] = resume.get("name","Unknown")
    result["filename"] = resume.get("filename","")
    result["experience_years"] = resume.get("experience_years",0)
    result["current_title"] = resume.get("current_title","")
    result["skills"] = resume.get("skills",[])
    result["internships"] = resume.get("internships",[])
    result["projects"] = resume.get("projects",[])
    result["raw_text"] = resume.get("raw_text","")   # store raw_text for interview auto-load
    return result

# ── AI: GENERATE INTERVIEW QUESTIONS (from JD tech stack) ─────────────────────
def generate_iv_questions(client, resume_text, jd_text, tech_stack):
    ts_note = f"\nTech stack required by JD: {tech_stack}" if tech_stack.strip() else ""
    prompt = f"""You are an expert technical interviewer. Generate exactly 7 interview questions:
- 2 about PROJECTS in the resume (category: "Project")
- 2 about tech stack / skills required in the JD (category: "Technical")
- 2 behavioral STAR-format (category: "Behavioral")
- 1 about career growth / experience (category: "Resume")
Return ONLY a valid JSON array, no markdown:
[{{"id":1,"question":"...","category":"Project|Technical|Behavioral|Resume","focus":"brief label"}}]
Resume:\"\"\"{resume_text[:3500]}\"\"\"{ts_note}
JD excerpt: {jd_text[:1000]}"""
    raw = call_groq(client, prompt)
    parsed = json.loads(clean_json(raw))
    qs = ensure_list(parsed)
    for j, q in enumerate(qs):
        if not isinstance(q, dict):
            qs[j] = {"id":j+1,"question":str(q),"category":"General","focus":""}
        else:
            q.setdefault("id",j+1); q.setdefault("category","General"); q.setdefault("focus","")
            if "question" not in q:
                q["question"] = f"Question {j+1}"
    return qs

# ── AI: INTERVIEW REPORT ──────────────────────────────────────────────────────

# Answers that count as "no response" — must score 0
EMPTY_ANSWER_MARKERS = [
    "[no response", "[skipped]", "[no response provided]",
    "[no response — time expired]", ""
]

def is_empty_answer(ans_text: str) -> bool:
    t = (ans_text or "").strip().lower()
    return any(t.startswith(m) or t == m for m in EMPTY_ANSWER_MARKERS) or len(t) < 5

def build_report_prompt(answers, resume_text):
    lines = []
    for i, a in enumerate(answers):
        ans = a['answer']
        empty_note = "  *** NO ANSWER PROVIDED — SCORE MUST BE 0 FOR THIS QUESTION ***" if is_empty_answer(ans) else ""
        lines.append(
            f"Q{i+1} [{a.get('category','')}]: {a['question']}\n"
            f"Answer ({a.get('word_count',0)}w, {a.get('time_taken','?')}s"
            f", {a.get('mode','?')}"
            f"{'  TIMED-OUT' if a.get('timed_out') else ''}"
            f"{empty_note}): {ans}"
        )
    qa_fmt = "\n\n".join(lines)

    return f"""You are a senior technical recruiter evaluating a software engineering candidate.
CRITICAL SCORING RULES — YOU MUST FOLLOW THESE EXACTLY:
1. If a candidate's answer is blank, "[No response", "[Skipped]", or fewer than 5 words, their score for that question MUST be 0. Do NOT infer from their resume.
2. Scores for overall_score, communication_score, technical_score, confidence_score, clarity_score must ALL be calculated ONLY from the actual answers given — NOT from the resume.
3. If most answers are empty, overall_score should reflect that (e.g. 0–20).
4. Return ONLY raw valid JSON — no markdown, no backticks, no explanation.
{{
  "candidate_name":"string","overall_score":0,"hire_confidence_pct":0,
  "recommendation":"Strong Hire|Hire|Consider|No Hire","recommendation_color":"green|blue|yellow|red",
  "executive_summary":"string","communication_score":0,"technical_score":0,"confidence_score":0,"clarity_score":0,
  "green_flags":["string"],"red_flags":["string"],"strengths":["string"],"areas_for_improvement":["string"],
  "question_evaluations":[{{"id":1,"question":"string","answer_summary":"string","score":0,
    "star_detected":false,"depth":"shallow|adequate|strong","specificity":"vague|moderate|specific",
    "feedback":"string","category":"string"}}],
  "final_verdict":"string"
}}
Resume (for context only — do NOT use to award scores): {resume_text[:2000]}
Interview Answers:
{qa_fmt}"""

def generate_iv_report(answers, resume_text):
    last_error = None
    for attempt in range(3):
        try:
            prompt = build_report_prompt(answers, resume_text)
            resp = st.session_state.groq_client.chat.completions.create(
                model=GROQ_MODEL,
                messages=[{"role":"user","content":prompt}],
                temperature=0.3, max_tokens=3500 if attempt==0 else 2500,
            )
            raw = resp.choices[0].message.content
            parsed = None
            for strategy in [
                lambda r: json.loads(clean_json(r)),
                lambda r: json.loads(re.sub(r',\s*([}\]])',r'\1',r[r.find('{'):r.rfind('}')+1])),
            ]:
                try:
                    parsed = strategy(raw); break
                except:
                    pass
            if parsed:
                parsed.setdefault("candidate_name","Candidate")
                parsed.setdefault("overall_score",0)
                parsed.setdefault("hire_confidence_pct",parsed.get("overall_score",0))
                parsed.setdefault("recommendation","Consider")
                parsed.setdefault("recommendation_color","yellow")
                parsed.setdefault("executive_summary","Evaluation complete.")
                parsed.setdefault("communication_score",0)
                parsed.setdefault("technical_score",0)
                parsed.setdefault("confidence_score",0)
                parsed.setdefault("clarity_score",0)
                parsed.setdefault("green_flags",[]); parsed.setdefault("red_flags",[])
                parsed.setdefault("strengths",[]); parsed.setdefault("areas_for_improvement",[])
                parsed.setdefault("final_verdict","See evaluation above.")
                parsed.setdefault("question_evaluations",[])
                for key in ["overall_score","hire_confidence_pct","communication_score","technical_score","confidence_score","clarity_score"]:
                    parsed[key] = safe_int(parsed.get(key, 0))

                # Hard-enforce: zero out scores for empty/skipped answers
                qe_list = parsed.get("question_evaluations", [])
                for i, qe in enumerate(qe_list):
                    orig_answer = answers[i]["answer"] if i < len(answers) else ""
                    if is_empty_answer(orig_answer):
                        qe["score"] = 0
                        qe["depth"] = "shallow"
                        qe["specificity"] = "vague"
                        qe["star_detected"] = False
                        qe["feedback"] = "No answer was provided for this question."
                        qe["answer_summary"] = "[No response]"
                    else:
                        qe["score"] = safe_int(qe.get("score", 0))
                        qe.setdefault("star_detected", False)
                        qe.setdefault("depth", "adequate")
                        qe.setdefault("specificity", "moderate")
                        qe.setdefault("feedback", "No feedback.")

                # Recalculate top-level scores proportionally based on answered questions
                n_total = len(answers)
                n_answered = sum(1 for a in answers if not is_empty_answer(a["answer"]))
                if n_answered == 0:
                    for key in ["overall_score","hire_confidence_pct","communication_score","technical_score","confidence_score","clarity_score"]:
                        parsed[key] = 0
                    parsed["recommendation"] = "No Hire"
                    parsed["recommendation_color"] = "red"
                    parsed["final_verdict"] = "Candidate did not answer any questions."
                elif n_answered < n_total:
                    ratio = n_answered / n_total
                    for key in ["overall_score","hire_confidence_pct","communication_score","technical_score","confidence_score","clarity_score"]:
                        parsed[key] = safe_int(round(parsed[key] * ratio))

                return parsed
        except Exception as e:
            last_error = e; time.sleep(1)
    raise RuntimeError(f"Report generation failed: {last_error}")

# ── HELPERS ────────────────────────────────────────────────────────────────────
def score_color(s):
    return "#34d399" if s>=75 else "#fbbf24" if s>=50 else "#f87171"

def score_class(s):
    return "score-high" if s>=75 else "score-mid" if s>=50 else "score-low"

def rec_class(v):
    return "rec-hire" if v=="Recommended" else "rec-maybe" if v=="Consider" else "rec-nope"

def rec_color_hex(c):
    return {"green":"#34d399","blue":"#60a5fa","yellow":"#fbbf24","red":"#f87171"}.get(c,"#60a5fa")

def score_bar(score, label):
    color = score_color(score)
    st.markdown(f"""<div style="margin-bottom:12px">
        <div style="display:flex;justify-content:space-between;margin-bottom:4px">
            <span style="font-size:0.82rem;color:#8090a0">{label}</span>
            <span style="font-size:0.82rem;font-family:'JetBrains Mono',monospace;color:{color}">{score}/100</span>
        </div>
        <div class="bar-wrap"><div class="bar-fill" style="width:{score}%;background:{color}"></div></div>
    </div>""", unsafe_allow_html=True)

def skill_tags(skills, css_class="skill-tag"):
    if skills:
        st.markdown("".join(f'<span class="{css_class}">{s}</span>' for s in skills[:15]), unsafe_allow_html=True)

def radar_chart(candidate_name, scores: dict):
    categories = list(scores.keys()); values = list(scores.values())
    values_closed = values + [values[0]]; cats_closed = categories + [categories[0]]
    fig = go.Figure()
    fig.add_trace(go.Scatterpolar(
        r=values_closed, theta=cats_closed, fill='toself',
        fillcolor='rgba(59,130,246,0.15)',
        line=dict(color='#60a5fa', width=2), name=candidate_name))
    fig.update_layout(
        polar=dict(bgcolor='#0d1422',
            radialaxis=dict(visible=True,range=[0,100],color='#3a5070',gridcolor='#1e2a45',tickfont=dict(color='#5a7090',size=10)),
            angularaxis=dict(color='#8090a0',gridcolor='#1e2a45',tickfont=dict(color='#b0c4de',size=11))),
        paper_bgcolor='#111827', plot_bgcolor='#111827',
        font=dict(color='#e8eaf0',family='Sora'),
        margin=dict(t=40,b=40,l=40,r=40), height=320, showlegend=False,
        title=dict(text=f"<b>{candidate_name}</b>",font=dict(color='#e8eaf0',size=14),x=0.5))
    return fig

# ── IV: SAVE ANSWER ────────────────────────────────────────────────────────────
def iv_save_answer(answer, timed_out=False):
    idx = st.session_state.iv_q_idx
    qs  = st.session_state.iv_questions
    q   = qs[idx]
    wc  = len(answer.split()) if answer else 0
    mode = "speech" if st.session_state.iv_live_transcript.strip() else "typed"
    st.session_state.iv_answers.append({
        "question": q["question"], "answer": answer,
        "category": q.get("category",""),
        "time_taken": int(time.time() - (st.session_state.iv_q_start_time or time.time())),
        "mode": mode, "timed_out": timed_out, "word_count": wc,
    })
    st.session_state.iv_q_idx += 1
    st.session_state.iv_q_start_time = time.time()
    st.session_state.iv_live_transcript = ""
    st.session_state.iv_typed_answer = ""
    st.session_state.iv_show_edit = False
    if st.session_state.iv_q_idx >= len(qs):
        st.session_state.iv_page = "report"

# ── HELPER: Get top resume raw_text from scored_results ───────────────────────
def get_top_resume_data():
    """Returns (name, filename, raw_text) of the #1 ranked candidate, or None."""
    if not st.session_state.scored_results:
        return None
    top = st.session_state.scored_results[0]
    name     = top.get("name","")
    filename = top.get("filename","")
    # raw_text is stored in scored_results (via score_resume) or fall back to parsed_resumes
    raw_text = top.get("raw_text","")
    if not raw_text:
        # try finding in parsed_resumes
        for r in st.session_state.parsed_resumes:
            if r.get("filename","") == filename:
                raw_text = r.get("raw_text","")
                break
    return {"name": name, "filename": filename, "raw_text": raw_text} if raw_text else None

# ══════════════════════════════════════════════════════════════════════════════
# SIDEBAR
# ══════════════════════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("## 🎯 TalentLens v3")
    st.markdown("<p style='color:#5a7090;font-size:0.82rem'>AI Hiring Suite · Resume + Interview</p>", unsafe_allow_html=True)
    st.markdown("---")

    # Show connection status (no API key input — it's hardcoded)
    if st.session_state.groq_client:
        st.success("✅ AI Connected")
    else:
        st.error("❌ AI connection failed — check API key in backend.")

    st.markdown("---")
    st.markdown("### ⚖️ Scoring Weights")
    st.caption("Must total 100%")
    w_skills  = st.slider("Skills",      10, 60, 30, 5)
    w_exp     = st.slider("Experience",  10, 60, 30, 5)
    w_edu     = st.slider("Education",   5,  40, 20, 5)
    w_culture = st.slider("Culture Fit", 5,  40, 20, 5)
    total_w   = w_skills + w_exp + w_edu + w_culture
    if total_w != 100:
        st.warning(f"⚠️ Total = {total_w}%")
    else:
        st.success(f"✅ Total = {total_w}%")
        st.session_state.weights = {"skills":w_skills,"experience":w_exp,"education":w_edu,"culture_fit":w_culture}

    st.markdown("---")
    st.markdown("### 📊 Session Stats")
    c1, c2 = st.columns(2)
    c1.metric("Resumes", len(st.session_state.parsed_resumes))
    c2.metric("Scored",  len(st.session_state.scored_results))
    if st.session_state.parsed_resumes:
        st.markdown("**Candidates:**")
        for r in st.session_state.parsed_resumes:
            st.markdown(f"<p style='font-size:0.8rem;color:#7090b0;margin:2px 0'>👤 {r.get('name','Unknown')}</p>", unsafe_allow_html=True)

    st.markdown("---")
    if st.button("🗑️ Reset Everything"):
        for k in ["parsed_resumes","scored_results","jd_text","jd_requirement","jd_tech_stack",
                  "interview_results","iv_page","iv_questions","iv_answers",
                  "iv_q_idx","iv_resume_text","iv_report","iv_live_transcript","iv_typed_answer",
                  "iv_auto_loaded"]:
            if k in st.session_state:
                val = st.session_state[k]
                st.session_state[k] = [] if isinstance(val,list) else ({} if isinstance(val,dict) else ("setup" if k=="iv_page" else (False if k=="iv_auto_loaded" else "")))
        st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
# HERO
# ══════════════════════════════════════════════════════════════════════════════
st.markdown("""
<div class="hero">
    <p class="hero-title">🎯 TalentLens v3</p>
    <p class="hero-sub">AI Hiring Suite · Score · Compare · Interview (Speech-Primary) · Report</p>
</div>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# TABS  (About removed)
# ══════════════════════════════════════════════════════════════════════════════
tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
    "📋 Job Description",
    "📂 Upload Resumes",
    "🏆 Score & Compare",
    "🎤 AI Interview",
    "📊 Interview Report",
    "📑 Final Report",
])

# ─────────────────────────────────────────────────────────────────────────────
# TAB 1 — JOB DESCRIPTION
# ─────────────────────────────────────────────────────────────────────────────
with tab1:
    st.markdown('<p class="section-title">Step 1 · Define the Role</p>', unsafe_allow_html=True)
    col_l, col_r = st.columns([1,1], gap="large")

    with col_l:
        st.markdown("#### ✍️ Describe what you're hiring for")
        requirement = st.text_area("Role", height=150, label_visibility="collapsed",
            placeholder="e.g. Python developer, 2+ years, FastAPI, AWS. Freshers with strong projects welcome.")
        if st.button("⚡ Generate Job Description", use_container_width=True):
            if not st.session_state.groq_client:
                st.error("❌ AI not connected.")
            elif not requirement.strip():
                st.error("❌ Describe the role.")
            else:
                with st.spinner("Writing JD..."):
                    try:
                        jd = generate_jd(st.session_state.groq_client, requirement)
                        st.session_state.jd_text = jd
                        st.session_state.jd_requirement = requirement
                        with st.spinner("🔍 Auto-extracting tech stack from JD..."):
                            tech = extract_tech_from_jd(st.session_state.groq_client, jd)
                            st.session_state.jd_tech_stack = tech
                        st.success("✅ JD ready! Tech stack auto-extracted.")
                    except Exception as e:
                        st.error(f"❌ {e}")

        st.markdown("#### 📝 Or paste your own JD")
        manual_jd = st.text_area("Paste JD", height=160, label_visibility="collapsed",
                                   placeholder="Paste an existing JD here...")
        if st.button("Use This JD", use_container_width=True):
            if manual_jd.strip():
                st.session_state.jd_text = manual_jd
                st.success("✅ JD saved! Click 'Extract Tech Stack' below.")
        if st.session_state.jd_text and st.button("🔍 Extract Tech Stack from JD", use_container_width=True):
            if not st.session_state.groq_client:
                st.error("❌ AI not connected.")
            else:
                with st.spinner("Extracting technologies..."):
                    try:
                        tech = extract_tech_from_jd(st.session_state.groq_client, st.session_state.jd_text)
                        st.session_state.jd_tech_stack = tech
                        st.success("✅ Tech stack extracted!")
                    except Exception as e:
                        st.error(f"❌ {e}")

    with col_r:
        if st.session_state.jd_text:
            st.markdown("#### 📄 Current Job Description")
            st.markdown(f'<div class="jd-box">{st.session_state.jd_text[:3500]}</div>', unsafe_allow_html=True)
            if st.session_state.jd_tech_stack:
                st.markdown("#### 🤖 Auto-Extracted Tech Stack")
                st.markdown(
                    '<div style="background:#0d1422;border:1px solid #09b394;border-radius:10px;padding:12px 16px;margin-bottom:12px">'
                    '<div style="font-family:\'JetBrains Mono\',monospace;font-size:10px;color:#09b394;letter-spacing:.1em;margin-bottom:8px">FROM JD · AUTO-DETECTED</div>'
                    + "".join(f'<span class="tech-tag-auto">{t.strip()}</span>'
                               for t in st.session_state.jd_tech_stack.split(',') if t.strip())
                    + "</div>",
                    unsafe_allow_html=True)
                st.caption("This tech stack is automatically passed to the AI Interview simulator.")
            st.download_button("📥 Download JD", st.session_state.jd_text,
                               file_name="job_description.txt", use_container_width=True)
        else:
            st.markdown("""<div style="background:#0d1422;border:1px dashed #1e3050;border-radius:12px;
                        padding:60px;text-align:center;color:#3a5070">
                <p style="font-size:2.5rem;margin:0">📄</p>
                <p style="margin:12px 0 0 0">Your JD will appear here</p>
            </div>""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# TAB 2 — UPLOAD RESUMES
# ─────────────────────────────────────────────────────────────────────────────
with tab2:
    st.markdown('<p class="section-title">Step 2 · Upload Candidate Resumes</p>', unsafe_allow_html=True)
    if not st.session_state.jd_text:
        st.warning("⚠️ Create a Job Description in Step 1 first.")
    uploaded = st.file_uploader("Drop PDF or DOCX files", type=["pdf","docx"], accept_multiple_files=True)
    if uploaded:
        st.info(f"📎 {len(uploaded)} file(s) ready")
        if st.button("🤖 Parse All Resumes with AI", use_container_width=True):
            if not st.session_state.groq_client:
                st.error("❌ AI not connected.")
            else:
                st.session_state.parsed_resumes = []
                progress = st.progress(0); status = st.empty()
                for i, file in enumerate(uploaded):
                    status.markdown(f"**Parsing** `{file.name}` ({i+1}/{len(uploaded)})...")
                    try:
                        text = read_pdf(file) if file.name.lower().endswith(".pdf") else read_docx(file)
                        parsed = parse_resume(st.session_state.groq_client, text, file.name)
                        st.session_state.parsed_resumes.append(parsed)
                    except Exception as e:
                        st.error(f"❌ {file.name}: {e}")
                    progress.progress((i+1)/len(uploaded)); time.sleep(1)
                status.success(f"✅ Parsed {len(st.session_state.parsed_resumes)} resumes!")
                st.balloons()

    if st.session_state.parsed_resumes:
        st.markdown('<p class="section-title">Parsed Candidate Profiles</p>', unsafe_allow_html=True)
        for r in st.session_state.parsed_resumes:
            with st.expander(f"👤  {r.get('name','Unknown')}  ·  {r.get('current_title','Fresher')}  ·  {r.get('experience_years',0)} yrs"):
                c1, c2, c3 = st.columns(3)
                with c1:
                    st.markdown("**📋 Basic Info**")
                    st.markdown(f"📧 {r.get('email') or 'N/A'}")
                    st.markdown(f"📞 {r.get('phone') or 'N/A'}")
                    st.markdown("**🎓 Education**")
                    for e in r.get("education",[])[:3]: st.markdown(f"- {e}")
                with c2:
                    st.markdown("**💼 Internships**")
                    for i in r.get("internships",[]):
                        st.markdown(f"**{i.get('role','')}** @ {i.get('company','')} ({i.get('duration','')})")
                        st.caption(i.get('description',''))
                    if not r.get("internships"): st.caption("No internships found")
                with c3:
                    st.markdown("**🚀 Projects**")
                    for p in r.get("projects",[])[:3]:
                        st.markdown(f"**{p.get('name','')}**"); st.caption(p.get('description',''))
                        skill_tags(p.get('tech_used',[]), "project-tag")
                    if not r.get("projects"): st.caption("No projects found")
                st.markdown("**🛠️ Skills & Technologies**")
                skill_tags(r.get("skills",[]) + r.get("technologies",[]))
                st.markdown(f"**Summary:** {r.get('summary','N/A')}")

# ─────────────────────────────────────────────────────────────────────────────
# TAB 3 — SCORE & COMPARE
# ─────────────────────────────────────────────────────────────────────────────
with tab3:
    st.markdown('<p class="section-title">Step 3 · Score & Compare Candidates</p>', unsafe_allow_html=True)
    if not st.session_state.parsed_resumes:
        st.warning("⚠️ Upload resumes in Step 2 first.")
    elif not st.session_state.jd_text:
        st.warning("⚠️ Create a JD in Step 1 first.")
    else:
        w = st.session_state.weights
        st.markdown(f"""<div style="background:#0d1422;border:1px solid #1e3050;border-radius:10px;padding:14px 20px;margin-bottom:20px;display:flex;gap:24px;flex-wrap:wrap">
            <span style="font-size:0.82rem;color:#5a7090">Active Weights:</span>
            <span style="font-size:0.82rem;color:#60a5fa">Skills {w['skills']}%</span>
            <span style="font-size:0.82rem;color:#a78bfa">Experience {w['experience']}%</span>
            <span style="font-size:0.82rem;color:#34d399">Education {w['education']}%</span>
            <span style="font-size:0.82rem;color:#fbbf24">Culture Fit {w['culture_fit']}%</span>
        </div>""", unsafe_allow_html=True)

        if total_w != 100:
            st.error("⚠️ Weights must total 100% before scoring.")
        else:
            if st.button("⚡ Score All Candidates", use_container_width=True):
                if not st.session_state.groq_client:
                    st.error("❌ AI not connected.")
                else:
                    st.session_state.scored_results = []
                    st.session_state.iv_auto_loaded = False   # reset auto-load flag on new scoring
                    progress = st.progress(0); status = st.empty()
                    for i, resume in enumerate(st.session_state.parsed_resumes):
                        status.markdown(f"**Scoring** `{resume.get('name','?')}` ({i+1}/{len(st.session_state.parsed_resumes)})...")
                        try:
                            score = score_resume(st.session_state.groq_client, resume, st.session_state.jd_text, st.session_state.weights)
                            st.session_state.scored_results.append(score)
                        except Exception as e:
                            st.error(f"❌ {resume.get('name','?')}: {e}")
                        progress.progress((i+1)/len(st.session_state.parsed_resumes)); time.sleep(1)
                    st.session_state.scored_results.sort(key=lambda x: x.get("overall_score",0), reverse=True)
                    status.success("✅ All candidates scored and ranked!")

        if st.session_state.scored_results:
            # ── Banner: top candidate auto-loaded for interview
            top_r = st.session_state.scored_results[0]
            st.markdown(f"""<div class="top-resume-banner">
                <div style="font-family:'JetBrains Mono',monospace;font-size:10px;letter-spacing:.1em;color:#34d399;margin-bottom:6px">🏆 TOP CANDIDATE · AUTO-LOADED FOR INTERVIEW</div>
                <div style="font-size:15px;font-weight:600;color:#e8eaf0">{top_r.get('name','?')}</div>
                <div style="font-size:12px;color:#5a7090;margin-top:3px">{top_r.get('current_title','N/A')} · Score: {top_r.get('overall_score',0)}/100 · {top_r.get('verdict','')}</div>
                <div style="font-size:12px;color:#09b394;margin-top:6px">→ This resume is automatically pre-loaded into the AI Interview tab. No upload needed.</div>
            </div>""", unsafe_allow_html=True)

            st.markdown("#### 🕸️ Radar Charts")
            cols = st.columns(min(len(st.session_state.scored_results), 3))
            for i, r in enumerate(st.session_state.scored_results):
                with cols[i % 3]:
                    fig = radar_chart(r.get("name","?"), {
                        "Skills": r.get("skills_score",0), "Experience": r.get("experience_score",0),
                        "Education": r.get("education_score",0), "Culture Fit": r.get("culture_fit_score",0),
                        "Projects": r.get("project_score",0), "Internships": r.get("internship_score",0),
                    }); st.plotly_chart(fig, use_container_width=True)

            st.markdown("#### 🏅 Ranked Candidates")
            for rank, r in enumerate(st.session_state.scored_results, 1):
                score = r.get("overall_score",0); verdict = r.get("verdict","Consider")
                crown = " 👑" if rank == 1 else ""
                st.markdown(f"""<div class="candidate-card">
                    <div style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:12px">
                        <div>
                            <span style="color:#5a7090;font-size:0.8rem;font-family:'JetBrains Mono',monospace">#{rank}</span>
                            <span class="candidate-name" style="margin-left:10px">{r.get('name','Unknown')}{crown}</span>
                            <div class="candidate-meta">{r.get('current_title','N/A')} · {r.get('experience_years',0)} yrs · {r.get('filename','')}</div>
                        </div>
                        <div style="display:flex;align-items:center;gap:12px">
                            <span class="score-pill {score_class(score)}">{score}/100</span>
                            <span class="rec-badge {rec_class(verdict)}">{verdict}</span>
                        </div>
                    </div>
                </div>""", unsafe_allow_html=True)
                with st.expander(f"Full breakdown — {r.get('name','?')}"):
                    ca, cb = st.columns(2)
                    with ca:
                        score_bar(r.get("skills_score",0),"Skills"); score_bar(r.get("experience_score",0),"Experience")
                        score_bar(r.get("education_score",0),"Education"); score_bar(r.get("culture_fit_score",0),"Culture Fit")
                        score_bar(r.get("internship_score",0),"Internships"); score_bar(r.get("project_score",0),"Projects")
                        score_bar(r.get("overall_score",0),"⭐ Overall (Weighted)")
                    with cb:
                        st.markdown("**✅ Strengths**")
                        for s in r.get("strengths",[]): st.markdown(f"- {s}")
                        st.markdown("**⚠️ Gaps**")
                        for g in r.get("gaps",[]): st.markdown(f"- {g}")
                    st.markdown(f"**💡 Verdict:** {r.get('recommendation','')} — {r.get('recommendation_reason','')}")

# ─────────────────────────────────────────────────────────────────────────────
# TAB 4 — AI INTERVIEW (Speech-Primary, top resume auto-loaded)
# ─────────────────────────────────────────────────────────────────────────────
with tab4:
    st.markdown('<p class="section-title">Step 4 · AI Interview Simulator</p>', unsafe_allow_html=True)

    # Handle URL query param actions for interview
    p = st.query_params
    action = p.get("action","")
    if action == "iv_transcript":
        new_tx = p.get("value","")
        if new_tx != st.session_state.iv_live_transcript:
            st.session_state.iv_live_transcript = new_tx
        st.query_params.clear(); st.rerun()
    elif action == "iv_timeout" and st.session_state.iv_page == "interview":
        idx = st.session_state.iv_q_idx
        if idx < len(st.session_state.iv_questions) and len(st.session_state.iv_answers) <= idx:
            answer = (st.session_state.iv_live_transcript.strip()
                      or st.session_state.iv_typed_answer.strip()
                      or p.get("answer","").strip()
                      or "[No response — time expired]")
            iv_save_answer(answer, timed_out=True)
        st.query_params.clear(); st.rerun()

    # ── SETUP PAGE ──────────────────────────────────────────────────────────
    if st.session_state.iv_page == "setup":
        st.markdown("""<div style="padding:14px 18px;background:rgba(13,242,200,.05);
            border:1px solid rgba(13,242,200,.2);border-radius:12px;margin-bottom:20px">
            <div style="font-family:'JetBrains Mono',monospace;font-size:10px;letter-spacing:.1em;color:#0df2c8;margin-bottom:8px">🎥 HOW IT WORKS</div>
            <div style="font-size:13px;color:#6a90b0;line-height:1.7">
            • 🏆 <strong style="color:#0df2c8">Top-ranked resume is automatically loaded</strong> — no manual upload needed<br>
            • 🤖 Tech stack is <strong style="color:#0df2c8">automatically extracted from your JD</strong><br>
            • 🔊 AI reads each question aloud automatically<br>
            • 🎙 Your speech is the <strong style="color:#0df2c8">primary answer</strong> — transcript shown live<br>
            • ✏️ Optional: click <strong style="color:#0df2c8">Edit Answer</strong> only if you want to correct<br>
            • ⏱ 60-second timer — auto-submits your spoken answer when it hits zero<br>
            • 📊 Generates a recruiter-grade report: STAR detection, red flags, hire confidence %
            </div></div>""", unsafe_allow_html=True)

        col1, col2 = st.columns([1,1], gap="large")

        # ── Auto-load top resume from scored results ──────────────────────
        top_data = get_top_resume_data()

        with col1:
            if top_data:
                st.markdown(f"""<div class="top-resume-banner">
                    <div style="font-family:'JetBrains Mono',monospace;font-size:10px;letter-spacing:.1em;color:#34d399;margin-bottom:8px">🏆 AUTO-LOADED · TOP CANDIDATE FROM SCORING</div>
                    <div style="display:flex;align-items:center;gap:12px">
                        <span style="font-size:28px">👤</span>
                        <div>
                            <div style="font-weight:600;font-size:16px;color:#e8eaf0">{top_data['name']}</div>
                            <div style="color:#5a7090;font-size:12px;margin-top:2px">{top_data['filename']}</div>
                            <div style="color:#34d399;font-size:11px;margin-top:4px">✓ Resume text ready · {len(top_data['raw_text'].split())} words extracted</div>
                        </div>
                    </div>
                </div>""", unsafe_allow_html=True)
                st.caption("The top-scoring candidate's resume was automatically loaded. You can also upload a different resume below.")
            else:
                st.info("⚠️ No scored candidates yet. Score candidates in Tab 3, or upload a resume manually below.")

            st.markdown('<div style="color:#5a7090;font-size:0.82rem;margin:10px 0 4px 0">📎 Or upload a different resume:</div>', unsafe_allow_html=True)
            iv_uploaded = st.file_uploader("Resume (PDF/DOCX/TXT)", type=["pdf","docx","txt"], key="iv_resume_upload")
            if iv_uploaded:
                st.markdown(f"""<div style="display:flex;align-items:center;gap:10px;padding:12px 16px;
                    background:rgba(13,242,200,.07);border:1px solid #09b394;border-radius:9px;margin-top:8px">
                    <span style="font-size:22px">📄</span>
                    <div><div style="font-weight:500;font-size:14px">{iv_uploaded.name}</div>
                    <div style="color:#6a90b0;font-size:12px">{round(iv_uploaded.size/1024,1)} KB</div></div>
                    <span style="color:#2ecc82;font-size:18px;margin-left:auto">✓ Will override auto-loaded resume</span></div>""",
                    unsafe_allow_html=True)

        with col2:
            st.markdown('<div class="ai-card-title">🤖 Auto-Detected Tech Stack from JD</div>', unsafe_allow_html=True)
            if st.session_state.jd_tech_stack:
                st.markdown(
                    '<div style="background:#0d1422;border:1px solid #09b394;border-radius:10px;padding:12px 14px;">'
                    + "".join(f'<span class="tech-tag-auto">{t.strip()}</span>'
                               for t in st.session_state.jd_tech_stack.split(',') if t.strip())
                    + '</div><p style="font-size:0.78rem;color:#5a7090;margin-top:6px">✅ Extracted from JD automatically</p>',
                    unsafe_allow_html=True)
            else:
                st.info("⚠️ No tech stack detected yet. Go to Tab 1 and generate/paste a JD first.")
                manual_tech = st.text_input("Or type tech stack manually:", placeholder="Python, React, AWS…", key="iv_manual_tech")
                if manual_tech:
                    st.session_state.jd_tech_stack = manual_tech

            if not st.session_state.jd_text:
                st.warning("⚠️ No JD found. Generate one in Step 1 first.")

        st.markdown("<br>", unsafe_allow_html=True)

        # Enable start if we have top_data OR a manual upload
        can_start = bool(top_data or iv_uploaded) and bool(st.session_state.groq_client)
        if st.button("🎤  Start AI Interview", use_container_width=True, disabled=not can_start):
            if not st.session_state.groq_client:
                st.error("❌ AI not connected.")
            else:
                with st.spinner("⬡ Analysing resume & generating personalised questions…"):
                    try:
                        # Prefer manual upload over auto-load
                        if iv_uploaded:
                            fb = iv_uploaded.read()
                            resume_text = extract_text_bytes(fb, iv_uploaded.name)
                            candidate_name = iv_uploaded.name
                        else:
                            resume_text = top_data["raw_text"]
                            candidate_name = top_data["name"]

                        if len(resume_text) < 50:
                            st.error("Could not extract text from resume."); st.stop()
                        qs = generate_iv_questions(
                            st.session_state.groq_client, resume_text,
                            st.session_state.jd_text or "",
                            st.session_state.jd_tech_stack or "")
                        if not qs:
                            st.error("Could not parse questions. Please try again."); st.stop()
                        st.session_state.update({
                            "iv_session_id": str(uuid.uuid4()),
                            "iv_questions":  qs, "iv_answers": [], "iv_q_idx": 0,
                            "iv_resume_text": resume_text,
                            "iv_candidate_name": candidate_name,
                            "iv_q_start_time": time.time(),
                            "iv_page": "interview",
                            "iv_live_transcript": "", "iv_typed_answer": "",
                            "iv_show_edit": False, "iv_report": None,
                            "iv_auto_loaded": True,
                        })
                        st.rerun()
                    except Exception as e:
                        st.error(f"Error: {e}")

    # ── INTERVIEW PAGE ──────────────────────────────────────────────────────
    elif st.session_state.iv_page == "interview":
        qs = st.session_state.iv_questions; idx = st.session_state.iv_q_idx; total = len(qs)
        if not isinstance(qs, list) or total == 0:
            st.error("No questions found. Please restart.")
            if st.button("↺ Back to Setup", key="iv_back_err"):
                st.session_state.iv_page = "setup"; st.rerun()
        elif idx >= total:
            st.session_state.iv_page = "report"; st.rerun()
        else:
            q = qs[idx]
            st.markdown(
                f'<div style="display:flex;justify-content:space-between;margin-bottom:6px">'
                f'<span style="font-family:\'JetBrains Mono\',monospace;font-size:11px;color:#6a90b0;letter-spacing:.08em">INTERVIEW IN PROGRESS</span>'
                f'<span style="font-family:\'Sora\',sans-serif;font-weight:700;color:#0df2c8;font-size:13px">Q{idx+1} / {total}</span></div>',
                unsafe_allow_html=True)
            st.progress(idx / total); st.markdown("<br>", unsafe_allow_html=True)

            left, right = st.columns([1, 1.7], gap="large")
            with left:
                cat_col = {"Project":"#34d399","Resume":"#0df2c8","Technical":"#fbbf24","Behavioral":"#60a5fa"}.get(q.get("category",""),"#0df2c8")
                focus   = q.get("focus","")
                st.markdown(f"""<div class="ai-card" style="text-align:center">
                    <div style="margin-bottom:16px">
                      <div style="width:90px;height:90px;border-radius:50%;background:linear-gradient(135deg,#0d2a42,#1e4a6e);
                        border:2px solid #0df2c8;display:inline-flex;align-items:center;justify-content:center;font-size:42px">🤖</div>
                    </div>
                    <div style="font-family:'JetBrains Mono',monospace;font-size:10px;letter-spacing:.1em;text-transform:uppercase;color:{cat_col};margin-bottom:12px">
                      ── {q.get('category','')} {('· '+focus) if focus else ''} ──</div>
                    <div class="q-bubble" style="text-align:left">{q['question']}</div>
                    </div>""", unsafe_allow_html=True)
                if st.session_state.iv_answers:
                    with st.expander(f"📋 Previous ({len(st.session_state.iv_answers)})", expanded=False):
                        for i, ans in enumerate(st.session_state.iv_answers):
                            pq = qs[i]
                            to_icon = "⏱ " if ans.get("timed_out") else ""
                            mode_icon = "🎙" if ans.get("mode")=="speech" else "✏️"
                            st.markdown(f"""<div style="margin-bottom:10px">
                                <div style="font-family:'JetBrains Mono',monospace;font-size:10px;color:#6a90b0">Q{i+1} · {pq.get('category','')} {to_icon}{mode_icon}</div>
                                <div style="font-size:12px;color:#e8eaf0;margin:2px 0">{pq['question'][:70]}…</div>
                                <div class="answer-bubble">{ans['answer'][:120]}{'…' if len(ans['answer'])>120 else ''}</div>
                            </div>""", unsafe_allow_html=True)

            with right:
                safe_q = (q['question'].replace('\\','\\\\').replace('"','\\"').replace('`',"'").replace('\n',' '))
                panel_html = fr"""<!DOCTYPE html><html><head><meta charset="utf-8">
<style>
*{{box-sizing:border-box;margin:0;padding:0;}}
body{{background:#060c14;font-family:sans-serif;color:#d8eaf8;}}
@keyframes pulse{{0%,100%{{opacity:1}}50%{{opacity:.25}}}}
@keyframes blink{{0%,100%{{opacity:1}}50%{{opacity:.5}}}}
@keyframes bar{{from{{opacity:.4;transform:scaleY(.4)}}to{{opacity:1;transform:scaleY(1)}}}}
#box{{position:relative;border-radius:14px;overflow:hidden;background:#060c14;border:2px solid #2a5a8a;}}
#vid{{width:100%;height:175px;object-fit:cover;display:block;background:#060c14;}}
#ov{{position:absolute;inset:0;display:flex;flex-direction:column;align-items:center;justify-content:center;
  background:rgba(6,12,20,.95);border-radius:12px;transition:opacity .6s;z-index:5;}}
.oi{{font-size:28px;margin-bottom:6px;}}
.om{{font-size:10px;letter-spacing:.1em;color:#0df2c8;text-align:center;padding:0 12px;}}
#tw{{position:absolute;top:8px;right:10px;z-index:10;display:none;}}
#rp{{position:absolute;top:8px;left:10px;z-index:10;display:none;
  background:rgba(255,78,106,.15);border:1px solid #ff4e6a;border-radius:20px;
  padding:3px 9px;font-size:9px;color:#ff4e6a;letter-spacing:.05em;}}
#rd{{display:inline-block;width:5px;height:5px;border-radius:50%;background:#ff4e6a;
  margin-right:4px;animation:pulse 1s infinite;vertical-align:middle;}}
#sb{{position:absolute;bottom:8px;left:10px;z-index:10;display:none;align-items:flex-end;gap:3px;height:18px;}}
.b{{width:3px;border-radius:2px;background:#0df2c8;animation:bar .8s ease-in-out infinite alternate;}}
.b:nth-child(1){{height:5px}}.b:nth-child(2){{height:12px;animation-delay:.15s}}
.b:nth-child(3){{height:8px;animation-delay:.3s}}.b:nth-child(4){{height:16px;animation-delay:.1s}}
.b:nth-child(5){{height:7px;animation-delay:.25s}}
#ab{{position:absolute;bottom:8px;right:10px;z-index:10;display:none;
  background:rgba(77,166,255,.18);border:1px solid #4da6ff;border-radius:20px;
  padding:3px 9px;font-size:9px;color:#4da6ff;animation:blink 1.2s infinite;}}
.txlabel{{display:flex;align-items:center;justify-content:space-between;
  margin:8px 0 4px;font-size:9px;letter-spacing:.1em;text-transform:uppercase;color:#6a90b0;}}
#ls{{font-size:9px;padding:2px 8px;border-radius:10px;background:rgba(13,242,200,.08);border:1px solid #09b394;color:#0df2c8;}}
#txbox{{background:#080f1a;border:2px solid rgba(13,242,200,.25);border-radius:12px;
  padding:12px 12px 30px;min-height:88px;max-height:115px;overflow-y:auto;position:relative;transition:border-color .3s,box-shadow .3s;}}
#txbox.active{{border-color:#0df2c8;box-shadow:0 0 14px rgba(13,242,200,.1);}}
#txtext{{font-size:14px;line-height:1.65;color:#d8eaf8;word-wrap:break-word;}}
#txtext.empty{{color:#2a4a6a;font-style:italic;font-size:12px;}}
#txfooter{{position:absolute;bottom:6px;left:12px;right:12px;display:flex;justify-content:space-between;align-items:center;}}
#txwc{{font-family:monospace;font-size:9px;color:#09b394;letter-spacing:.06em;}}
#txst{{font-family:monospace;font-size:9px;color:#0df2c8;letter-spacing:.06em;}}
#interim{{color:#4da6ff;font-style:italic;}}
#nb{{display:none;padding:5px 10px;border-radius:8px;font-size:11px;text-align:center;
  background:rgba(245,166,35,.08);border:1px solid #f5a623;color:#f5a623;margin-top:6px;}}
#tb{{display:none;padding:5px 10px;border-radius:8px;font-size:11px;text-align:center;
  background:rgba(255,78,106,.08);border:1px solid #ff4e6a;color:#ff4e6a;margin-top:6px;}}
</style></head><body>
<div id="box">
  <video id="vid" autoplay muted playsinline></video>
  <div id="ov"><div class="oi" id="oi">🔊</div><div class="om" id="om">AI IS READING…</div></div>
  <div id="tw"><svg width="48" height="48" viewBox="0 0 60 60">
    <circle cx="30" cy="30" r="24" fill="none" stroke="#1a3550" stroke-width="4"/>
    <circle id="ta" cx="30" cy="30" r="24" fill="none" stroke="#0df2c8" stroke-width="4"
      stroke-linecap="round" stroke-dasharray="150.8" stroke-dashoffset="0"
      style="transform:rotate(-90deg);transform-origin:30px 30px;transition:stroke-dashoffset .95s linear,stroke .3s"/>
    <text id="tn" x="30" y="35" text-anchor="middle"
      style="fill:#d8eaf8;font-size:14px;font-family:Arial,sans-serif;font-weight:800">60</text>
  </svg></div>
  <div id="rp"><span id="rd"></span>REC</div>
  <div id="sb"><div class="b"></div><div class="b"></div><div class="b"></div><div class="b"></div><div class="b"></div></div>
  <div id="ab">🔊 AI Speaking…</div>
</div>
<div class="txlabel">🎙 YOUR ANSWER — LIVE SPEECH (PRIMARY) <span id="ls">WAITING…</span></div>
<div id="txbox">
  <div id="txtext" class="empty">Speak after the question — your words appear here live…</div>
  <div id="txfooter"><span id="txwc"></span><span id="txst"></span></div>
</div>
<div id="nb">⚠️ No speech detected — speak louder or use Edit box below</div>
<div id="tb">⏱ Time's up — spoken answer auto-submitted!</div>
<script>
(function(){{
  var TOTAL=60,ARC=150.8;
  var stream=null,recog=null,tmr=null,alertTO=null,syncTO=null;
  var secs=TOTAL,spoken=false,finished=false;
  var finalTx='',interimTx='',aiSpk=false,lastSynced='';
  var vid=document.getElementById('vid'),ov=document.getElementById('ov'),
      oi=document.getElementById('oi'),om=document.getElementById('om'),
      tw=document.getElementById('tw'),ta=document.getElementById('ta'),tn=document.getElementById('tn'),
      rp=document.getElementById('rp'),sb=document.getElementById('sb'),ab=document.getElementById('ab'),
      ls=document.getElementById('ls'),txbox=document.getElementById('txbox'),
      txtext=document.getElementById('txtext'),txwc=document.getElementById('txwc'),
      txst=document.getElementById('txst'),nb=document.getElementById('nb'),tb=document.getElementById('tb');
  function signal(params){{
    try{{
      var qs=Object.keys(params).map(function(k){{return encodeURIComponent(k)+'='+encodeURIComponent(params[k]);}}).join('&');
      window.parent.history.replaceState(null,'','?'+qs);
      window.parent.dispatchEvent(new Event('popstate'));
    }}catch(e){{console.warn('signal err',e);}}
  }}
  function schedSync(){{
    if(syncTO)clearTimeout(syncTO);
    syncTO=setTimeout(function(){{
      var tx=(finalTx+interimTx).trim();
      if(tx&&tx!==lastSynced){{lastSynced=tx;signal({{action:'iv_transcript',value:tx}});}}
    }},1500);
  }}
  function updatePanel(){{
    var full=finalTx.trim(),interim=interimTx.trim();
    if(full||interim){{
      txtext.classList.remove('empty');
      txtext.innerHTML=full+(interim?' <span id="interim">'+interim+'</span>':'');
      txbox.classList.add('active');
      var wc=(full+' '+interim).trim().split(/[\s]+/).filter(Boolean).length;
      txwc.textContent=wc+' words'; txst.textContent='● SPEECH CAPTURED';
      txbox.scrollTop=txbox.scrollHeight;
    }}else{{
      txtext.classList.add('empty');
      txtext.textContent='Speak after the question — your words appear here live…';
      txbox.classList.remove('active'); txwc.textContent=''; txst.textContent='';
    }}
  }}
  function autoSubmit(){{
    finished=true; stopAll(); tb.style.display='block';
    var tx=(finalTx+interimTx).trim();
    signal({{action:'iv_timeout',answer:tx}});
  }}
  function stopAll(){{
    clearInterval(tmr); clearTimeout(alertTO); clearTimeout(syncTO);
    if(window.speechSynthesis)window.speechSynthesis.cancel();
    if(recog)try{{recog.stop();}}catch(e){{}}
    if(stream)stream.getTracks().forEach(function(t){{t.stop();}});
    rp.style.display='none'; sb.style.display='none';
    ls.textContent='DONE'; ls.style.color='#6a90b0';
  }}
  function speak(text,cb){{
    if(!window.speechSynthesis){{cb();return;}}
    window.speechSynthesis.cancel();
    var u=new SpeechSynthesisUtterance(text);
    u.rate=0.92; u.lang='en-US';
    var vs=window.speechSynthesis.getVoices();
    var v=vs.find(function(x){{return x.lang==='en-US'&&(x.name.includes('Google')||x.name.includes('Samantha'));}})
         ||vs.find(function(x){{return x.lang.startsWith('en');}});
    if(v)u.voice=v;
    aiSpk=true; ab.style.display='block';
    u.onend=u.onerror=function(){{aiSpk=false;ab.style.display='none';cb();}};
    window.speechSynthesis.speak(u);
  }}
  function initTTS(){{
    var vs=window.speechSynthesis.getVoices();
    if(vs.length>0)speak("{safe_q}",openCam);
    else window.speechSynthesis.onvoiceschanged=function(){{speak("{safe_q}",openCam);}};
  }}
  function openCam(){{
    oi.textContent='📷'; om.textContent='REQUESTING CAMERA…';
    navigator.mediaDevices.getUserMedia({{video:true,audio:true}})
      .then(function(s){{
        stream=s; vid.srcObject=s;
        vid.onloadedmetadata=function(){{
          ov.style.opacity='0';
          setTimeout(function(){{ov.style.display='none';}},600);
          tw.style.display='block'; rp.style.display='block';
          startSpeech(); startTimer();
          alertTO=setTimeout(function(){{if(!spoken)nb.style.display='block';}},12000);
        }};
      }}).catch(function(){{
        ov.innerHTML='<div style="color:#ff4e6a;font-size:12px;text-align:center;padding:20px">⚠️ Camera blocked<br><span style="color:#6a90b0;font-size:11px">Mic-only mode</span></div>';
        tw.style.display='block'; startTimer();
        navigator.mediaDevices.getUserMedia({{audio:true}})
          .then(function(s){{stream=s;startSpeech();}}).catch(function(){{}});
      }});
  }}
  function startSpeech(){{
    var SR=window.SpeechRecognition||window.webkitSpeechRecognition;
    if(!SR){{ls.textContent='USE EDIT BOX'; txst.textContent='⌨ TYPE MODE'; return;}}
    recog=new SR();
    recog.continuous=true; recog.interimResults=true; recog.lang='en-US';
    recog.onresult=function(e){{
      if(aiSpk)return;
      var interim='';
      for(var i=e.resultIndex;i<e.results.length;i++){{
        if(e.results[i].isFinal)finalTx+=e.results[i][0].transcript+' ';
        else interim=e.results[i][0].transcript;
      }}
      interimTx=interim; updatePanel();
      ls.textContent='● REC'; ls.style.color='#ff4e6a';
      if(!spoken){{spoken=true;nb.style.display='none';sb.style.display='flex';}}
      schedSync();
    }};
    recog.onerror=function(e){{if(e.error!=='no-speech')ls.textContent='MIC ERR';}};
    recog.onend=function(){{
      sb.style.display='none'; ls.textContent='LISTENING'; ls.style.color='#0df2c8';
      if(!finished)try{{recog.start();}}catch(x){{}}
    }};
    recog.start(); ls.textContent='LISTENING'; ls.style.color='#0df2c8';
  }}
  function startTimer(){{
    secs=TOTAL; drawArc(TOTAL);
    tmr=setInterval(function(){{secs--;drawArc(secs);if(secs<=0){{clearInterval(tmr);autoSubmit();}}}},1000);
  }}
  function drawArc(s){{
    ta.style.strokeDashoffset=ARC*(1-s/TOTAL); tn.textContent=Math.max(0,s);
    ta.style.stroke=s<=10?'#ff4e6a':s<=20?'#f5a623':'#0df2c8';
    tn.style.fill=s<=10?'#ff4e6a':'#d8eaf8';
  }}
  if(window.speechSynthesis)initTTS(); else openCam();
}})();
</script></body></html>"""
                components.html(panel_html, height=420, scrolling=False)

                tx = st.session_state.iv_live_transcript
                if tx:
                    wc = len(tx.split())
                    preview = tx[:90]+("…" if len(tx)>90 else "")
                    st.markdown(f'<div class="tx-pill">🎙 SPEECH READY · {wc} words · "{preview}"</div>', unsafe_allow_html=True)
                else:
                    st.markdown('<div style="font-size:11px;color:#3a5a7a;font-family:monospace;margin:4px 0">🎙 Speak above — transcript syncs here automatically</div>', unsafe_allow_html=True)

                show_edit = st.session_state.get("iv_show_edit", False)
                edit_col, _ = st.columns([1,2])
                with edit_col:
                    if not show_edit:
                        if st.button("✏️  Edit Answer", key=f"iv_showedit_{idx}", use_container_width=True):
                            st.session_state.iv_show_edit = True; st.rerun()
                    else:
                        if st.button("✕  Close Edit", key=f"iv_hideedit_{idx}", use_container_width=True):
                            st.session_state.iv_show_edit = False; st.rerun()

                edit_value = ""
                if show_edit:
                    st.markdown('<div style="font-family:\'JetBrains Mono\',monospace;font-size:9px;letter-spacing:.08em;text-transform:uppercase;color:#f5a623;margin:6px 0 3px">✏️ EDIT OVERRIDE</div>', unsafe_allow_html=True)
                    def _on_iv_type():
                        st.session_state.iv_typed_answer = st.session_state.get(f"iv_ta_{idx}","")
                    edit_value = st.text_area("Edit Answer", value=st.session_state.iv_typed_answer or tx,
                        height=100, placeholder="Edit your spoken answer or type from scratch…",
                        label_visibility="collapsed", key=f"iv_ta_{idx}", on_change=_on_iv_type)

                b1, b2 = st.columns([3,1])
                with b1:
                    if st.button("✓  Submit & Next Question", use_container_width=True, type="primary", key=f"iv_sub_{idx}"):
                        typed  = st.session_state.iv_typed_answer.strip()
                        widget = (edit_value or "").strip() if show_edit else ""
                        speech = tx.strip()
                        final  = (typed or widget) if (typed or widget) else speech
                        if not final: final = "[No response provided]"
                        iv_save_answer(final, timed_out=False); st.rerun()
                with b2:
                    if st.button("Skip →", use_container_width=True, key=f"iv_skp_{idx}"):
                        iv_save_answer("[Skipped]", timed_out=False); st.rerun()

    elif st.session_state.iv_page == "report":
        st.info("✅ Interview complete! Head to Tab 5 — Interview Report to view your results.")
        if st.button("↺ New Interview", key="iv_new_from_done"):
            st.session_state.iv_page = "setup"
            st.session_state.iv_questions = []; st.session_state.iv_answers = []
            st.session_state.iv_q_idx = 0; st.session_state.iv_report = None
            st.rerun()

# ─────────────────────────────────────────────────────────────────────────────
# TAB 5 — INTERVIEW REPORT
# ─────────────────────────────────────────────────────────────────────────────
with tab5:
    st.markdown('<p class="section-title">Interview Evaluation Report</p>', unsafe_allow_html=True)

    if st.session_state.iv_page != "report" and not st.session_state.iv_report:
        if not st.session_state.iv_answers:
            st.warning("⚠️ Complete the AI Interview in Step 4 first.")
        else:
            st.info("Interview in progress — complete all questions to generate the report.")
    else:
        if not st.session_state.iv_report and st.session_state.iv_answers:
            prog = st.progress(0); status = st.empty()
            try:
                status.markdown("""<div style="display:flex;align-items:center;gap:12px;padding:16px 20px;
                  background:rgba(13,242,200,.05);border:1px solid rgba(13,242,200,.2);border-radius:12px">
                  <div style="font-size:24px">⬡</div>
                  <div><div style="font-family:'JetBrains Mono',monospace;font-size:10px;letter-spacing:.1em;color:#0df2c8">GENERATING REPORT</div>
                  <div style="font-size:13px;color:#6a90b0;margin-top:4px">Analysing answers with recruiter-grade evaluation…</div></div>
                </div>""", unsafe_allow_html=True)
                prog.progress(20)
                report = generate_iv_report(st.session_state.iv_answers, st.session_state.iv_resume_text)
                prog.progress(100); st.session_state.iv_report = report
                status.empty(); prog.empty(); st.rerun()
            except Exception as e:
                prog.empty(); status.empty()
                st.error(f"Report generation failed: {e}")
                if st.button("↺ Try Again", key="iv_retry_report"):
                    st.session_state.iv_report = None; st.rerun()

        if st.session_state.iv_report:
            r = st.session_state.iv_report
            rc = rec_color_hex(r.get("recommendation_color","blue"))
            ov = safe_int(r.get("overall_score",0))
            hc = safe_int(r.get("hire_confidence_pct",ov))

            st.markdown(f"""<div style="background:linear-gradient(135deg,rgba(13,242,200,.06),rgba(77,166,255,.06));
              border:1px solid #1a3550;border-radius:16px;padding:24px 28px;margin-bottom:28px;
              display:flex;align-items:center;justify-content:space-between;flex-wrap:wrap;gap:16px">
              <div>
                <div style="font-family:'JetBrains Mono',monospace;font-size:10px;letter-spacing:.15em;text-transform:uppercase;color:#6a90b0;margin-bottom:6px">⬡ AI Interview · Recruiter Report</div>
                <div style="font-family:'Sora',sans-serif;font-size:28px;font-weight:700;line-height:1.1">{r.get('candidate_name','Candidate')}</div>
                <div style="font-family:'JetBrains Mono',monospace;font-size:11px;color:#3a5a7a;margin-top:4px">{datetime.now().strftime('%B %d, %Y · %H:%M')}</div>
              </div>
              <div style="text-align:center">
                <div style="font-size:11px;font-family:'JetBrains Mono',monospace;color:#6a90b0;margin-bottom:6px">RECOMMENDATION</div>
                <div style="font-family:'Sora',sans-serif;font-weight:700;font-size:20px;color:{rc};padding:8px 20px;
                  border:2px solid {rc};border-radius:10px;background:rgba(0,0,0,.3)">{r.get('recommendation','N/A')}</div>
              </div>
            </div>""", unsafe_allow_html=True)

            m1, m2, m3, m4 = st.columns(4)
            for col, label in zip([m1,m2,m3,m4],[
                ("Overall Score", f"{ov}/100"),("Hire Confidence", f"{hc}%"),
                ("Communication", f"{safe_int(r.get('communication_score',0))}/100"),
                ("Technical", f"{safe_int(r.get('technical_score',0))}/100")]):
                with col:
                    st.metric(label[0], label[1])

            st.markdown("<br>", unsafe_allow_html=True)
            s1, s2 = st.columns(2)
            with s1:
                scores = [("Communication",safe_int(r.get('communication_score',0))),
                          ("Technical",safe_int(r.get('technical_score',0))),
                          ("Confidence",safe_int(r.get('confidence_score',0))),
                          ("Clarity",safe_int(r.get('clarity_score',0)))]
                bars = ""
                for lbl, v in scores:
                    c = score_color(v)
                    bars += f"""<div style="margin-bottom:12px">
                      <div style="display:flex;justify-content:space-between;font-size:12px;margin-bottom:5px">
                        <span style="color:#6a90b0">{lbl}</span>
                        <span style="color:{c};font-family:'JetBrains Mono',monospace;font-size:11px;font-weight:700">{v}/100</span>
                      </div>
                      <div style="height:6px;background:#1a3550;border-radius:3px;overflow:hidden">
                        <div style="height:6px;background:{c};border-radius:3px;width:{v}%"></div>
                      </div></div>"""
                st.markdown(f'<div class="ai-card"><div class="ai-card-title">📊 Skill Scores</div>{bars}</div>', unsafe_allow_html=True)
            with s2:
                st.markdown(f"""<div class="ai-card" style="height:100%">
                    <div class="ai-card-title">📋 Executive Summary</div>
                    <p style="color:#9ab8d0;font-size:13.5px;line-height:1.75;margin:0">{r.get('executive_summary','')}</p>
                </div>""", unsafe_allow_html=True)

            gfc, rfc = st.columns(2)
            default_text = "<span style='color:#3a5a7a'>None</span>"
            with gfc:
                flags = r.get("green_flags",[])
                items = "".join(f'<span class="flag-green">✦ {f}</span>' for f in flags)
                st.markdown(f'<div class="ai-card"><div class="ai-card-title" style="color:#34d399">✅ Green Flags</div>{items or default_text}</div>', unsafe_allow_html=True)
            with rfc:
                flags = r.get("red_flags",[])
                items = "".join(f'<span class="flag-red">⚠ {f}</span>' for f in flags)
                st.markdown(f'<div class="ai-card"><div class="ai-card-title" style="color:#f87171">🚩 Red Flags</div>{items or default_text}</div>', unsafe_allow_html=True)

            st.markdown('<p class="section-title">Question-by-Question Evaluation</p>', unsafe_allow_html=True)
            for i, qa in enumerate(r.get("question_evaluations",[])):
                sv = safe_int(qa.get("score",0))
                orig = st.session_state.iv_answers[i] if i < len(st.session_state.iv_answers) else {}
                mode_icon = "🎙" if orig.get("mode")=="speech" else "✏️"
                star = "⭐ " if qa.get("star_detected") else ""
                with st.expander(f"Q{i+1} · {qa.get('category','')} · {mode_icon} · Score: {sv}/100 {star}"):
                    ca, cb = st.columns([2,1])
                    with ca:
                        st.markdown(f"**{qa.get('question','')}**")
                        stored = orig.get("answer", qa.get("answer_summary",""))
                        st.markdown(f'<div class="answer-bubble">{stored}</div>', unsafe_allow_html=True)
                        if qa.get("feedback"):
                            st.markdown(f'<div class="feedback-bubble">💬 {qa["feedback"]}</div>', unsafe_allow_html=True)
                    with cb:
                        score_bar(sv, "Score")
                        st.markdown(f"**Depth:** {qa.get('depth','—')}")
                        st.markdown(f"**Specificity:** {qa.get('specificity','—')}")

            rc2 = rec_color_hex(r.get("recommendation_color","blue"))
            st.markdown(f"""<div class="verdict-box">
              <div style="font-family:'JetBrains Mono',monospace;font-size:10px;letter-spacing:.12em;text-transform:uppercase;color:#0df2c8;margin-bottom:8px">⬡ Final Verdict</div>
              <div style="display:flex;gap:16px;align-items:center;flex-wrap:wrap;margin-bottom:12px">
                <div style="padding:8px 18px;border:1px solid {rc2};border-radius:10px;">
                  <div style="font-size:9px;color:#6a90b0;font-family:'JetBrains Mono',monospace">VERDICT</div>
                  <div style="font-weight:700;font-size:16px;color:{rc2}">{r.get('recommendation','N/A')}</div>
                </div>
                <div style="padding:8px 18px;border:1px solid {score_color(hc)};border-radius:10px;">
                  <div style="font-size:9px;color:#6a90b0;font-family:'JetBrains Mono',monospace">CONFIDENCE</div>
                  <div style="font-weight:700;font-size:16px;color:{score_color(hc)}">{hc}%</div>
                </div>
              </div>
              <p style="font-size:14px;line-height:1.8;color:#b0c4de;margin:0;border-top:1px solid #1e2a45;padding-top:12px">{r.get('final_verdict','')}</p>
            </div>""", unsafe_allow_html=True)

            st.markdown("<br>", unsafe_allow_html=True)
            ba, bb = st.columns(2)
            with ba:
                if st.button("↺ New Interview", use_container_width=True, key="iv_new_from_report"):
                    st.session_state.iv_page = "setup"
                    st.session_state.iv_questions = []; st.session_state.iv_answers = []
                    st.session_state.iv_q_idx = 0; st.session_state.iv_report = None
                    st.rerun()
            with bb:
                st.download_button("⬇ Download Report (JSON)", data=json.dumps(r, indent=2),
                    file_name=f"interview_report_{datetime.now().strftime('%Y%m%d_%H%M')}.json",
                    mime="application/json", use_container_width=True)

# ─────────────────────────────────────────────────────────────────────────────
# TAB 6 — FINAL HIRING REPORT  (Shortlist email section removed)
# ─────────────────────────────────────────────────────────────────────────────
with tab6:
    st.markdown('<p class="section-title">Final Hiring Report</p>', unsafe_allow_html=True)
    if not st.session_state.scored_results:
        st.warning("⚠️ Score candidates in Step 3 first.")
    else:
        results = st.session_state.scored_results
        total = len(results)
        recommended = sum(1 for r in results if r.get("verdict")=="Recommended")
        consider = sum(1 for r in results if r.get("verdict")=="Consider")
        not_rec = sum(1 for r in results if r.get("verdict")=="Not Recommended")
        avg_score = int(np.mean([r.get("overall_score",0) for r in results]))
        top = results[0]

        st.markdown(f"""<div class="metric-row">
            <div class="metric-card"><div class="metric-num">{total}</div><div class="metric-label">Evaluated</div></div>
            <div class="metric-card"><div class="metric-num" style="color:#34d399">{recommended}</div><div class="metric-label">Recommended</div></div>
            <div class="metric-card"><div class="metric-num" style="color:#fbbf24">{consider}</div><div class="metric-label">Consider</div></div>
            <div class="metric-card"><div class="metric-num" style="color:#f87171">{not_rec}</div><div class="metric-label">Not Recommended</div></div>
            <div class="metric-card"><div class="metric-num">{avg_score}</div><div class="metric-label">Avg Score</div></div>
            <div class="metric-card"><div class="metric-num" style="color:#a78bfa;font-size:1rem">{top.get('name','—')[:14]}</div><div class="metric-label">Top Pick</div></div>
        </div>""", unsafe_allow_html=True)

        st.markdown("#### 📋 Comparison Matrix")
        headers = ["Candidate","Overall","Skills","Exp.","Education","Culture","Projects","Internships","Verdict"]
        cols = st.columns([2,1,1,1,1,1,1,1,1.5])
        for col, h in zip(cols, headers):
            col.markdown(f"<p style='font-size:0.7rem;color:#5a7090;font-weight:600;text-transform:uppercase;letter-spacing:1px;margin:0'>{h}</p>", unsafe_allow_html=True)
        st.markdown("<hr style='margin:6px 0'>", unsafe_allow_html=True)
        for r in results:
            cols = st.columns([2,1,1,1,1,1,1,1,1.5])
            cols[0].markdown(f"<p style='font-size:0.85rem;color:#e8eaf0;margin:4px 0'>{r.get('name','?')}</p>", unsafe_allow_html=True)
            for ci, key in enumerate(["overall_score","skills_score","experience_score","education_score","culture_fit_score","project_score","internship_score"]):
                s = r.get(key,0)
                cols[ci+1].markdown(f"<p style='font-family:JetBrains Mono;font-size:0.85rem;color:{score_color(s)};margin:4px 0'>{s}</p>", unsafe_allow_html=True)
            v = r.get("verdict","Consider")
            cols[8].markdown(f'<span class="rec-badge {rec_class(v)}" style="font-size:0.7rem;padding:3px 8px">{v}</span>', unsafe_allow_html=True)

        st.markdown("---")
        lines = ["="*65,"    TALENTLENS v3 — AI HIRING SUITE REPORT","="*65,""]
        if st.session_state.jd_requirement:
            lines += [f"ROLE: {st.session_state.jd_requirement}",""]
        if st.session_state.jd_tech_stack:
            lines += [f"TECH STACK: {st.session_state.jd_tech_stack}",""]
        w = st.session_state.weights
        lines += [f"WEIGHTS: Skills {w['skills']}% | Experience {w['experience']}% | Education {w['education']}% | Culture {w['culture_fit']}%",
                  f"Total: {total} | Recommended: {recommended} | Consider: {consider} | Avg: {avg_score}/100","","-"*65,""]
        for rank, r in enumerate(results, 1):
            lines += [f"RANK #{rank}  —  {r.get('name','?')}",
                      f"  Overall: {r.get('overall_score',0)}/100 | Skills: {r.get('skills_score',0)} | Exp: {r.get('experience_score',0)} | Edu: {r.get('education_score',0)}",
                      f"  Verdict: {r.get('recommendation','')} — {r.get('recommendation_reason','')}",
                      "  Strengths:"]
            for s in r.get("strengths",[]): lines.append(f"    + {s}")
            lines.append("  Gaps:")
            for g in r.get("gaps",[]): lines.append(f"    - {g}")
            lines += ["","-"*65,""]
        st.download_button("📥 Download Full Report (.txt)", "\n".join(lines), file_name="talentlens_report.txt", mime="text/plain", use_container_width=True)

# ── FOOTER ─────────────────────────────────────────────────────────────────────
st.markdown("""
<div style="text-align:center;padding:32px 0 16px 0;color:#2a3a55;font-size:0.8rem">
    TalentLens v3 · AI Hiring Suite · Resume Evaluator + Speech Interview · Groq LLaMA 3.3
</div>""", unsafe_allow_html=True)